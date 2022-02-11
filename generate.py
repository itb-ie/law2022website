import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from contract_template import *
import os


def generate(app, lawyer, client, service, compensation_value, deposit_value, deposit_date, jurisdiction):
    # get all the information from the entry text and generate the file
    refundable_deposit_value = str(int(deposit_value)//2) if deposit_value else "0"
    nonrefundable_deposit_value = str(int(deposit_value)//2) if deposit_value else "0"

    # start generating the document
    doc = docx.Document()
    # title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.font.bold = True

    # parties
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(parties % (lawyer, client))
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # service
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(services % service)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # responsibilities
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(responsibilities)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # compensation
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(compensation % compensation_value)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(compensation2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(compensation3)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(compensation4)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # costs
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(costs)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # deposit
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(deposit % (deposit_value, deposit_date, refundable_deposit_value, nonrefundable_deposit_value))
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # provisions
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(provisions % jurisdiction)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # effective date
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(effective_date)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # foregoing
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(foregoing)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # signatures
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(signatures)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    filename = f"{client}-Contract.docx"

    # save the doc
    doc.save(os.path.join(app.config['CONTRACT_FOLDER'], filename))

    return filename
